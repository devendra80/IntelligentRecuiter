from docx import Document

def convertDocxToText(fPath):
    document = Document(fPath)
    content = "\n".join([p.text for p in document.paragraphs])
    tabText = ""
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                tabText = tabText + " " + cell.text
    #print('tabText :', "\n".join(content, tabText))            
    #return "\n".join(content, tabText)
    return content + "\n" + tabText

